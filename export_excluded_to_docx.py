#!/usr/bin/env python3
"""Combine all excluded Fragestellungen into a single .docx file.

Reads every .txt in <input_dir> (default:
C:/Users/kilia/Desktop/main_data_filtered/excluded), sorts by filename, and
writes them to a Word document with the case_id as a small heading and
exactly two line breaks (one blank paragraph) between cases.

Usage:
    uv run export_excluded_to_docx.py
    uv run export_excluded_to_docx.py --input <dir> --output <file.docx>
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.shared import Pt


DEFAULT_INPUT  = Path(r"C:\Users\kilia\Desktop\main_data_filtered\excluded")
DEFAULT_OUTPUT = Path(r"C:\Users\kilia\Desktop\main_data_filtered\excluded.docx")


def read_text(path: Path) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return path.read_text(encoding=enc).strip()
        except UnicodeDecodeError:
            continue
    return ""


def build_doc(input_dir: Path, output: Path) -> int:
    files = sorted(p for p in input_dir.iterdir()
                   if p.is_file() and p.suffix.lower() == ".txt")
    if not files:
        raise SystemExit(f"No .txt files in {input_dir}")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for i, fp in enumerate(files):
        if i > 0:
            # exactly two line breaks between cases = one empty paragraph
            doc.add_paragraph("")

        head = doc.add_paragraph()
        run = head.add_run(fp.stem)
        run.bold = True

        text = read_text(fp) or "(leer)"
        doc.add_paragraph(text)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return len(files)


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--input", type=Path, default=DEFAULT_INPUT,
                    help=f"Folder with excluded .txt files "
                         f"(default: {DEFAULT_INPUT}).")
    ap.add_argument("--output", type=Path, default=DEFAULT_OUTPUT,
                    help=f"Output .docx path (default: {DEFAULT_OUTPUT}).")
    args = ap.parse_args()

    if not args.input.is_dir():
        raise SystemExit(f"ERROR: input folder does not exist: {args.input}")

    n = build_doc(args.input, args.output)
    print(f"Wrote {n} cases to {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
